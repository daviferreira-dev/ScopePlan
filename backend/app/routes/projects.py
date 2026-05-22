from flask import Blueprint, request, jsonify, send_file
from flask_jwt_extended import jwt_required, get_jwt_identity
from app.models import Project, User, Requirement, AuditLog
from app.schemas import ProjectCreateSchema, ProjectUpdateSchema
from app.utils.decorators import validate_json, role_required
from app.utils.access import check_user_project_access
from app import db
import io

projects_bp = Blueprint('projects', __name__, url_prefix='/api/projects')


@projects_bp.route('', methods=['POST'])
@jwt_required()
@validate_json(ProjectCreateSchema)
def create_project():
    data = request.validated_data
    user_id = get_jwt_identity()
    user = db.session.get(User, user_id)

    if not user:
        return {'message': 'Usuário não encontrado'}, 404

    if user.perfil not in ('analista', 'gestor'):
        return {'message': 'Apenas analistas ou gestores podem criar projetos'}, 403

    if data.get('custo_estimado') is not None:
        data['custo_estimado'] = float(data['custo_estimado'])

    # Resolve cliente_id -> nome_cliente automatically
    cliente_id = data.get('cliente_id')
    nome_cliente = data.get('nome_cliente')
    if cliente_id:
        cliente = db.session.get(User, cliente_id)
        if cliente and cliente.perfil == 'cliente':
            nome_cliente = cliente.nome
        else:
            return {'message': 'Cliente não encontrado ou usuário não é cliente'}, 400

    project = Project(
        nome=data['nome'],
        descricao=data.get('descricao'),
        status=data.get('status', 'planejamento'),
        custo_estimado=data.get('custo_estimado'),
        gestor_id=user_id,
        cliente_id=cliente_id,
        nome_cliente=nome_cliente
    )
    db.session.add(project)
    db.session.commit()

    AuditLog.log(user_id, 'criacao', 'projeto', project.id, project.id, {'nome': project.nome})
    db.session.commit()

    return {'message': 'Projeto criado com sucesso', 'project': project.to_dict()}, 201


@projects_bp.route('', methods=['GET'])
@jwt_required()
def list_projects():
    user_id = get_jwt_identity()
    user = db.session.get(User, user_id)

    if not user:
        return {'message': 'Usuário não encontrado'}, 404

    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 20, type=int)
    per_page = min(per_page, 100)

    query = Project.query.filter_by(ativo=True)

    if user.perfil == 'cliente':
        query = query.filter_by(cliente_id=user.id)
    elif user.perfil in ('analista', 'gestor'):
        query = query.filter_by(gestor_id=user.id)
    elif user.perfil == 'desenvolvedor':
        project_ids = db.session.query(Requirement.projeto_id).filter_by(
            autor_id=user.id, ativo=True
        ).distinct().scalar_subquery()
        query = query.filter(Project.id.in_(project_ids))
    else:
        query = query.filter(Project.id == -1)  # no results

    pagination = query.order_by(Project.criado_em.desc()).paginate(
        page=page, per_page=per_page, error_out=False
    )

    return jsonify({
        'projetos': [p.to_dict() for p in pagination.items],
        'total': pagination.total,
        'page': pagination.page,
        'per_page': pagination.per_page,
        'pages': pagination.pages
    }), 200


@projects_bp.route('/<int:project_id>', methods=['GET'])
@jwt_required()
def get_project(project_id):
    user_id = get_jwt_identity()
    user = db.session.get(User, user_id)

    if not user:
        return {'message': 'Usuário não encontrado'}, 404

    project, access_error = check_user_project_access(user, project_id)
    if access_error:
        return access_error

    return jsonify(project.to_dict()), 200


@projects_bp.route('/<int:project_id>', methods=['PUT'])
@jwt_required()
@validate_json(ProjectUpdateSchema)
def update_project(project_id):
    data = request.validated_data
    user_id = get_jwt_identity()
    user = db.session.get(User, user_id)

    if not user:
        return {'message': 'Usuário não encontrado'}, 404

    if user.perfil not in ('analista', 'gestor'):
        return {'message': 'Apenas analistas ou gestores podem editar projetos'}, 403

    project, access_error = check_user_project_access(user, project_id)
    if access_error:
        return access_error

    # Additional check: only the gestor of this project can edit it
    if project.gestor_id != user.id:
        return {'message': 'Acesso negado a este projeto'}, 403

    # Update fields
    if 'nome' in data:
        project.nome = data['nome']
    if 'descricao' in data:
        project.descricao = data['descricao']
    if 'status' in data:
        project.status = data['status']
    if 'custo_estimado' in data:
        project.custo_estimado = data['custo_estimado']

    # Handle cliente_id change
    if 'cliente_id' in data:
        cliente_id = data['cliente_id']
        if cliente_id:
            cliente = db.session.get(User, cliente_id)
            if cliente and cliente.perfil == 'cliente':
                project.cliente_id = cliente_id
                project.nome_cliente = cliente.nome
            else:
                return {'message': 'Cliente não encontrado ou usuário não é cliente'}, 400

    db.session.commit()

    AuditLog.log(user_id, 'atualizacao', 'projeto', project.id, project.id, {'nome': project.nome})
    db.session.commit()

    return {'message': 'Projeto atualizado com sucesso', 'project': project.to_dict()}, 200


@projects_bp.route('/<int:project_id>', methods=['DELETE'])
@jwt_required()
def delete_project(project_id):
    user_id = get_jwt_identity()
    user = db.session.get(User, user_id)

    if not user:
        return {'message': 'Usuário não encontrado'}, 404

    if user.perfil not in ('analista', 'gestor'):
        return {'message': 'Apenas analistas ou gestores podem excluir projetos'}, 403

    project, access_error = check_user_project_access(user, project_id)
    if access_error:
        return access_error

    if project.gestor_id != user.id:
        return {'message': 'Acesso negado a este projeto'}, 403

    # Soft delete
    project.ativo = False
    db.session.commit()

    AuditLog.log(user_id, 'exclusao', 'projeto', project.id, project.id, {'nome': project.nome})
    db.session.commit()

    return {'message': 'Projeto excluído com sucesso'}, 200


@projects_bp.route('/<int:project_id>/ers/download', methods=['POST'])
@jwt_required()
def download_ers(project_id):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    user_id = get_jwt_identity()
    user = db.session.get(User, user_id)

    if not user:
        return {'message': 'Usuário não encontrado'}, 404

    project, access_error = check_user_project_access(user, project_id)
    if access_error:
        return access_error

    # Get format and topic_ids from request
    data = request.get_json() or {}
    formato = data.get('formato', request.args.get('format', 'docx'))
    topic_ids = data.get('topic_ids')  # Can be None or list

    # Get approved requirements only (RN002)
    reqs = Requirement.query.filter_by(projeto_id=project_id, ativo=True, status='aprovado').all()

    # Filter by topic_ids if provided
    if topic_ids and isinstance(topic_ids, list) and len(topic_ids) > 0:
        reqs = [r for r in reqs if r.id in topic_ids]

    # Sort by type and code — keys match model's tipo values
    type_order = {'funcional': 0, 'nao_funcional': 1, 'negocio': 2, 'restricao': 3}
    reqs.sort(key=lambda r: (type_order.get(r.tipo, 99), r.codigo))

    # Create DOCX
    doc = Document()

    # Title
    title = doc.add_heading('Especificação de Requisitos de Software', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f'Projeto: {project.nome}')
    doc.add_paragraph(f'Descrição: {project.descricao or "N/A"}')
    doc.add_paragraph(f'Cliente: {project.nome_cliente or "N/A"}')
    doc.add_paragraph(f'Data: {project.atualizado_em.strftime("%d/%m/%Y") if project.atualizado_em else "N/A"}')
    doc.add_paragraph()

    # Group by type — keys match model's tipo values
    type_names = {
        'funcional': 'Requisitos Funcionais',
        'nao_funcional': 'Requisitos Não-Funcionais',
        'negocio': 'Regras de Negócio',
        'restricao': 'Restrições'
    }

    for req_type in ['funcional', 'nao_funcional', 'negocio', 'restricao']:
        type_reqs = [r for r in reqs if r.tipo == req_type]
        if type_reqs:
            doc.add_heading(type_names.get(req_type, req_type), level=1)
            for req in type_reqs:
                p = doc.add_paragraph()
                p.add_run(f'{req.codigo} ').bold = True
                p.add_run(req.titulo)
                doc.add_paragraph(req.descricao or '')

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    AuditLog.log(user_id, 'download', 'requisito', project_id, project_id, {'formato': formato, 'topic_ids': topic_ids})
    db.session.commit()

    return send_file(
        buffer,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f'ers_{project.nome.replace(" ", "_").lower()}.docx'
    )

"""ERS (Especificação de Requisitos de Software) document generator.

Generates PDF and DOCX files from project requirements.
"""

import io
from datetime import datetime, timezone
from zoneinfo import ZoneInfo

BRT = ZoneInfo('America/Sao_Paulo')
def now_brt(): return datetime.now(BRT)

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
import xml.etree.ElementTree as ET


# ── Labels ────────────────────────────────────────────────────────────────────

TOPIC_LABELS = {
    'funcional':    'Requisitos Funcionais',
    'nao_funcional':'Requisitos Não Funcionais',
    'negocio':      'Regras de Negócio',
    'restricao':    'Restrições',
}

TOPIC_ORDER = ['funcional', 'nao_funcional', 'negocio', 'restricao']

PRIORITY_LABELS = {
    'baixa':  'Baixa',
    'media':  'Média',
    'alta':   'Alta',
    'critica':'Crítica',
}

STATUS_LABELS = {
    'rascunho':              'Rascunho',
    'em_revisao':            'Em Revisão',
    'aprovado':              'Aprovado',
    'aprovado_com_ressalvas':'Aprovado c/ Ressalvas',
    'rejeitado':             'Rejeitado',
}

STATUS_COLORS = {
    'aprovado':              (22,  163, 74),
    'aprovado_com_ressalvas':(234, 179, 8),
    'rejeitado':             (220, 38,  38),
    'em_revisao':            (59,  130, 246),
    'rascunho':              (107, 114, 128),
}

PRIORITY_COLORS = {
    'critica':(220, 38,  38),
    'alta':   (234, 179, 8),
    'media':  (59,  130, 246),
    'baixa':  (107, 114, 128),
}

# ScopePlan green
GREEN      = (26,  102, 52)
GREEN_LIGHT= (34,  197, 90)
DARK       = (15,  23,  42)
GRAY_BG    = (248, 250, 252)
GRAY_BORDER= (203, 213, 225)


# ── Helpers ───────────────────────────────────────────────────────────────────

def _sort_reqs(reqs):
    """Sort requirements by numeric suffix in codigo (RF1 < RF2 < RF10), fallback to id."""
    import re
    def _key(r):
        codigo = r.get('codigo') or ''
        m = re.search(r'(\d+)$', codigo)
        return (int(m.group(1)) if m else 10**9, r.get('id', 0))
    return sorted(reqs, key=_key)


def _group_by_topic(requirements):
    groups = {}
    for req in requirements:
        tipo = req.get('tipo', 'funcional') or 'funcional'
        groups.setdefault(tipo, []).append(req)
    ordered = []
    for tipo in TOPIC_ORDER:
        if tipo in groups:
            ordered.append((tipo, _sort_reqs(groups[tipo])))
    for tipo in groups:
        if tipo not in TOPIC_ORDER:
            ordered.append((tipo, _sort_reqs(groups[tipo])))
    return ordered


def _filter_requirements(requirements, topic_ids=None, requirement_ids=None):
    filtered = requirements
    if requirement_ids:
        rid_set = set(requirement_ids)
        filtered = [r for r in filtered if r['id'] in rid_set]
    if topic_ids:
        filtered = [r for r in filtered if r.get('tipo') in topic_ids]
    return filtered


def _rgb(r, g, b):
    from reportlab.lib.colors import Color
    return Color(r / 255, g / 255, b / 255)


# ── PDF Builder ───────────────────────────────────────────────────────────────

def _build_pdf(project, requirements, topic_ids=None, requirement_ids=None,
               diagramas=None, custom_labels=None):
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm, mm
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
            PageBreak, HRFlowable, KeepTogether,
        )
        from reportlab.platypus import Image as RLImage
        from reportlab.lib import colors
        from reportlab.pdfgen import canvas as rl_canvas
        from xml.sax.saxutils import escape as xe
    except ImportError:
        return _build_minimal_pdf(project, requirements, topic_ids, requirement_ids)

    filtered = _filter_requirements(requirements, topic_ids, requirement_ids)
    groups   = _group_by_topic(filtered)
    clabels  = custom_labels or {}

    # ── Color palette ────────────────────────────────────────────────────────
    c_green       = _rgb(*GREEN)
    c_green_light = _rgb(*GREEN_LIGHT)
    c_dark        = _rgb(*DARK)
    c_gray_bg     = _rgb(*GRAY_BG)
    c_gray_border = _rgb(*GRAY_BORDER)
    c_white       = colors.white

    PAGE_W, PAGE_H = A4
    ML = MR = 2.2*cm
    MT = MB = 2.2*cm

    # ── Canvas callback for header/footer ────────────────────────────────────
    project_nome = project.get('nome', 'Projeto')

    def _on_page(canvas, doc):
        canvas.saveState()
        # Header bar
        canvas.setFillColor(c_green)
        canvas.rect(0, PAGE_H - 1.1*cm, PAGE_W, 1.1*cm, fill=1, stroke=0)
        canvas.setFont('Helvetica-Bold', 8)
        canvas.setFillColor(c_white)
        canvas.drawString(ML, PAGE_H - 0.75*cm, 'ERS — Especificação de Requisitos de Software')
        canvas.setFont('Helvetica', 8)
        canvas.drawRightString(PAGE_W - MR, PAGE_H - 0.75*cm, project_nome)
        # Footer
        canvas.setFillColor(c_gray_border)
        canvas.rect(ML, MB - 0.6*cm, PAGE_W - ML - MR, 0.5, fill=1, stroke=0)
        canvas.setFont('Helvetica', 7.5)
        canvas.setFillColor(_rgb(100, 116, 139))
        canvas.drawString(ML, MB - 0.55*cm,
                          f'Gerado em {now_brt().strftime("%d/%m/%Y às %H:%M")} (Horário de Brasília)  ·  ScopePlan')
        canvas.drawRightString(PAGE_W - MR, MB - 0.55*cm, f'Página {doc.page}')
        canvas.restoreState()

    def _on_first_page(canvas, doc):
        # Capa — fundo escuro
        canvas.saveState()
        canvas.setFillColor(c_dark)
        canvas.rect(0, 0, PAGE_W, PAGE_H, fill=1, stroke=0)
        # Barra lateral verde
        canvas.setFillColor(c_green)
        canvas.rect(0, 0, 0.7*cm, PAGE_H, fill=1, stroke=0)
        canvas.restoreState()

    # ── Styles ───────────────────────────────────────────────────────────────
    base = getSampleStyleSheet()

    def ps(name, **kw):
        return ParagraphStyle(name, **kw)

    # Cover styles
    cover_label = ps('CoverLabel', fontName='Helvetica', fontSize=9,
                     textColor=c_green_light, spaceAfter=6, alignment=TA_LEFT)
    cover_title = ps('CoverTitle', fontName='Helvetica-Bold', fontSize=30,
                     textColor=c_white, leading=38, spaceAfter=12, alignment=TA_LEFT)
    cover_sub   = ps('CoverSub',   fontName='Helvetica', fontSize=13,
                     textColor=_rgb(148,163,184), spaceAfter=4, alignment=TA_LEFT)
    cover_meta  = ps('CoverMeta',  fontName='Helvetica', fontSize=10,
                     textColor=_rgb(148,163,184), spaceAfter=3, alignment=TA_LEFT)

    # Body styles
    h1 = ps('ERSH1', fontName='Helvetica-Bold', fontSize=15, textColor=c_green,
             spaceBefore=18, spaceAfter=8, leading=20,
             borderPad=0, leftIndent=0)
    h2 = ps('ERSH2', fontName='Helvetica-Bold', fontSize=11.5, textColor=c_dark,
             spaceBefore=14, spaceAfter=5, leading=16)
    body = ps('ERSBody', fontName='Helvetica', fontSize=10, textColor=c_dark,
              leading=15, spaceAfter=6, alignment=TA_JUSTIFY)
    meta_label = ps('MetaLabel', fontName='Helvetica-Bold', fontSize=8.5,
                    textColor=_rgb(71,85,105))
    meta_value = ps('MetaValue', fontName='Helvetica', fontSize=8.5,
                    textColor=c_dark)
    toc_entry  = ps('TOCEntry', fontName='Helvetica', fontSize=10,
                    textColor=c_dark, leading=18)
    toc_title  = ps('TOCTitle', fontName='Helvetica-Bold', fontSize=13,
                    textColor=c_green, spaceAfter=10)
    caption    = ps('Caption', fontName='Helvetica-Oblique', fontSize=8.5,
                    textColor=_rgb(100,116,139), alignment=TA_CENTER, spaceAfter=8)
    section_intro = ps('SectionIntro', fontName='Helvetica', fontSize=9.5,
                       textColor=_rgb(100,116,139), spaceAfter=10, alignment=TA_JUSTIFY)

    elements = []

    # ═══════════════════════════════════════════════════════════════════════════
    # CAPA
    # ═══════════════════════════════════════════════════════════════════════════
    gestor      = project.get('gestor') or {}
    nome_gestor = gestor.get('nome', 'N/A')
    nome_cliente= project.get('nome_cliente') or '—'
    data_hoje   = now_brt().strftime('%d/%m/%Y')

    # Empurra conteúdo para o meio da página
    elements.append(Spacer(1, 5.5*cm))
    elements.append(Paragraph('ESPECIFICAÇÃO DE REQUISITOS DE SOFTWARE', cover_label))
    elements.append(Paragraph(xe(project_nome), cover_title))
    elements.append(Spacer(1, 0.4*cm))
    elements.append(HRFlowable(width='100%', thickness=1, color=c_green, spaceAfter=16))
    elements.append(Paragraph(f'Cliente: {xe(nome_cliente)}',  cover_sub))
    elements.append(Paragraph(f'Gestor: {xe(nome_gestor)}',   cover_sub))
    elements.append(Paragraph(f'Data: {data_hoje}',           cover_sub))
    elements.append(Spacer(1, 1*cm))
    total_reqs = len(filtered)
    total_aprv = sum(1 for r in filtered if r.get('status') in ('aprovado','aprovado_com_ressalvas'))
    stats_data = [['Requisitos', 'Aprovados', 'Seções']]
    stats_data.append([str(total_reqs), str(total_aprv), str(len(groups))])
    stats_t = Table(stats_data, colWidths=[4*cm, 4*cm, 4*cm])
    stats_t.setStyle(TableStyle([
        ('BACKGROUND',  (0,0),(-1,0), _rgb(*GREEN)),
        ('TEXTCOLOR',   (0,0),(-1,0), c_white),
        ('FONTNAME',    (0,0),(-1,0), 'Helvetica-Bold'),
        ('FONTSIZE',    (0,0),(-1,0), 9),
        ('BACKGROUND',  (0,1),(-1,1), _rgb(30,41,59)),
        ('TEXTCOLOR',   (0,1),(-1,1), c_white),
        ('FONTNAME',    (0,1),(-1,1), 'Helvetica-Bold'),
        ('FONTSIZE',    (0,1),(-1,1), 18),
        ('ALIGN',       (0,0),(-1,-1),'CENTER'),
        ('VALIGN',      (0,0),(-1,-1),'MIDDLE'),
        ('TOPPADDING',  (0,0),(-1,-1), 8),
        ('BOTTOMPADDING',(0,0),(-1,-1), 8),
        ('GRID',        (0,0),(-1,-1), 0.5, _rgb(*GREEN)),
        ('ROUNDEDCORNERS',(0,0),(-1,-1), 4),
    ]))
    elements.append(stats_t)
    elements.append(PageBreak())

    # ═══════════════════════════════════════════════════════════════════════════
    # SUMÁRIO
    # ═══════════════════════════════════════════════════════════════════════════
    elements.append(Paragraph('Sumário', toc_title))
    elements.append(HRFlowable(width='100%', thickness=1.5, color=c_green, spaceAfter=14))

    toc_items = []
    sec = 1
    if project.get('descricao'):
        toc_items.append((f'{sec}. Descrição do Projeto', ''))
        sec += 1
    for tipo, reqs in groups:
        label = TOPIC_LABELS.get(tipo, clabels.get(tipo, tipo))
        toc_items.append((f'{sec}. {label}', str(len(reqs))))
        sec += 1
    if diagramas:
        toc_items.append((f'{sec}. Diagramas', str(len(diagramas))))
        sec += 1
    toc_items.append((f'{sec}. Resumo', ''))

    for item_text, count in toc_items:
        row_data = [[Paragraph(item_text, toc_entry),
                     Paragraph(f'{count} item(s)' if count else '', meta_value)]]
        toc_row = Table(row_data, colWidths=[13*cm, 3*cm])
        toc_row.setStyle(TableStyle([
            ('ALIGN',        (1,0),(1,0),'RIGHT'),
            ('VALIGN',       (0,0),(-1,-1),'MIDDLE'),
            ('BOTTOMPADDING',(0,0),(-1,-1), 6),
            ('TOPPADDING',   (0,0),(-1,-1), 2),
            ('LINEBELOW',    (0,0),(-1,0), 0.3, c_gray_border),
        ]))
        elements.append(toc_row)

    elements.append(PageBreak())

    # ═══════════════════════════════════════════════════════════════════════════
    # DESCRIÇÃO DO PROJETO
    # ═══════════════════════════════════════════════════════════════════════════
    section_num = 1
    if project.get('descricao'):
        elements.append(Paragraph(f'{section_num}. Descrição do Projeto', h1))
        elements.append(HRFlowable(width='100%', thickness=1, color=c_green_light, spaceAfter=10))
        elements.append(Paragraph(xe(project['descricao']), body))
        elements.append(Spacer(1, 0.5*cm))
        section_num += 1

    # ═══════════════════════════════════════════════════════════════════════════
    # SEÇÕES DE REQUISITOS
    # ═══════════════════════════════════════════════════════════════════════════
    for tipo, reqs in groups:
        label = TOPIC_LABELS.get(tipo, clabels.get(tipo, tipo))

        elements.append(Paragraph(f'{section_num}. {label}', h1))
        elements.append(HRFlowable(width='100%', thickness=1, color=c_green_light, spaceAfter=6))
        elements.append(Paragraph(
            f'Esta seção contém {len(reqs)} item(s) do tipo <i>{label}</i>.', section_intro))

        for idx, req in enumerate(reqs, 1):
            codigo  = req.get('codigo') or f'{section_num}.{idx:03d}'
            titulo  = req.get('titulo', '')
            desc    = req.get('descricao', '') or ''
            prio    = req.get('prioridade') or 'media'
            status  = req.get('status') or 'rascunho'
            categ   = req.get('categoria') or '—'

            prio_color   = _rgb(*PRIORITY_COLORS.get(prio,   (107,114,128)))
            status_color = _rgb(*STATUS_COLORS.get(status, (107,114,128)))

            # Card header
            header_data = [[
                Paragraph(f'<font color="#{"%02x%02x%02x" % GREEN}">{xe(codigo)}</font>  '
                          f'<b>{xe(titulo)}</b>', h2),
            ]]
            card_header = Table(header_data, colWidths=[PAGE_W - ML - MR - 0.4*cm])
            card_header.setStyle(TableStyle([
                ('BACKGROUND',   (0,0),(-1,-1), c_gray_bg),
                ('LEFTPADDING',  (0,0),(-1,-1), 10),
                ('RIGHTPADDING', (0,0),(-1,-1), 10),
                ('TOPPADDING',   (0,0),(-1,-1), 8),
                ('BOTTOMPADDING',(0,0),(-1,-1), 6),
                ('LINEABOVE',    (0,0),(-1,0), 2, c_green),
                ('ROUNDEDCORNERS',(0,0),(-1,-1), 3),
            ]))

            # Description
            desc_data = [[Paragraph(xe(desc) if desc else '<i>Sem descrição.</i>', body)]]
            desc_table = Table(desc_data, colWidths=[PAGE_W - ML - MR - 0.4*cm])
            desc_table.setStyle(TableStyle([
                ('BACKGROUND',   (0,0),(-1,-1), c_white),
                ('LEFTPADDING',  (0,0),(-1,-1), 10),
                ('RIGHTPADDING', (0,0),(-1,-1), 10),
                ('TOPPADDING',   (0,0),(-1,-1), 6),
                ('BOTTOMPADDING',(0,0),(-1,-1), 6),
                ('LINEBEFORE',   (0,0),(0,-1), 2, c_green_light),
            ]))

            # Metadata chips row
            meta_data = [[
                Paragraph(f'<b>Prioridade:</b> {PRIORITY_LABELS.get(prio, prio)}', meta_label),
                Paragraph(f'<b>Status:</b> {STATUS_LABELS.get(status, status)}', meta_label),
            ]]
            meta_table = Table(meta_data, colWidths=[5*cm, 5.5*cm])
            meta_table.setStyle(TableStyle([
                ('BACKGROUND',   (0,0),(-1,-1), _rgb(241,245,249)),
                ('LEFTPADDING',  (0,0),(-1,-1), 10),
                ('RIGHTPADDING', (0,0),(-1,-1), 10),
                ('TOPPADDING',   (0,0),(-1,-1), 5),
                ('BOTTOMPADDING',(0,0),(-1,-1), 5),
                ('LINEBELOW',    (0,0),(-1,-1), 0.5, c_gray_border),
                ('TEXTCOLOR',    (0,0),(0,-1), prio_color),
                ('TEXTCOLOR',    (1,0),(1,-1), status_color),
            ]))

            elements.append(KeepTogether([card_header, desc_table, meta_table, Spacer(1, 0.3*cm)]))

        section_num += 1
        elements.append(Spacer(1, 0.3*cm))

    # ═══════════════════════════════════════════════════════════════════════════
    # DIAGRAMAS
    # ═══════════════════════════════════════════════════════════════════════════
    if diagramas:
        elements.append(Paragraph(f'{section_num}. Diagramas', h1))
        elements.append(HRFlowable(width='100%', thickness=1, color=c_green_light, spaceAfter=10))

        max_img_w = PAGE_W - ML - MR - 0.4*cm

        for diagrama in diagramas:
            nome_dia = diagrama.get('nome', 'Diagrama')
            dados    = diagrama.get('dados')
            tipo_mime= diagrama.get('tipo_mime', '')

            elements.append(Paragraph(xe(nome_dia), h2))

            if dados and 'pdf' not in tipo_mime:
                try:
                    img_buf = io.BytesIO(dados)
                    img = RLImage(img_buf)
                    # Scale proportionally to fit page width
                    scale = min(max_img_w / img.drawWidth, (PAGE_H * 0.6) / img.drawHeight)
                    img.drawWidth  *= scale
                    img.drawHeight *= scale
                    # Border frame
                    frame_data = [[img]]
                    frame = Table(frame_data, colWidths=[max_img_w])
                    frame.setStyle(TableStyle([
                        ('ALIGN',        (0,0),(-1,-1),'CENTER'),
                        ('VALIGN',       (0,0),(-1,-1),'MIDDLE'),
                        ('BOX',          (0,0),(-1,-1), 1, c_gray_border),
                        ('BACKGROUND',   (0,0),(-1,-1), c_white),
                        ('TOPPADDING',   (0,0),(-1,-1), 8),
                        ('BOTTOMPADDING',(0,0),(-1,-1), 8),
                    ]))
                    elements.append(frame)
                    elements.append(Paragraph(nome_dia, caption))
                except Exception:
                    elements.append(Paragraph('[Imagem não pôde ser inserida]', body))
            elif 'pdf' in tipo_mime:
                elements.append(Paragraph('[Arquivo PDF — visualize separadamente]', body))

            elements.append(Spacer(1, 0.5*cm))

        section_num += 1

    # ═══════════════════════════════════════════════════════════════════════════
    # RESUMO
    # ═══════════════════════════════════════════════════════════════════════════
    elements.append(Paragraph(f'{section_num}. Resumo', h1))
    elements.append(HRFlowable(width='100%', thickness=1, color=c_green_light, spaceAfter=10))

    summary_rows = [['Seção', 'Quantidade']]
    for tipo, reqs in groups:
        label = TOPIC_LABELS.get(tipo, clabels.get(tipo, tipo))
        summary_rows.append([label, str(len(reqs))])
    summary_rows.append(['Total', str(len(filtered))])

    summary_t = Table(summary_rows, colWidths=[13*cm, 3*cm])
    summary_t.setStyle(TableStyle([
        ('BACKGROUND',   (0,0),(-1,0), _rgb(*GREEN)),
        ('TEXTCOLOR',    (0,0),(-1,0), c_white),
        ('FONTNAME',     (0,0),(-1,0), 'Helvetica-Bold'),
        ('FONTSIZE',     (0,0),(-1,-1), 10),
        ('ALIGN',        (1,0),(1,-1), 'CENTER'),
        ('BACKGROUND',   (0,-1),(-1,-1), c_gray_bg),
        ('FONTNAME',     (0,-1),(-1,-1), 'Helvetica-Bold'),
        ('GRID',         (0,0),(-1,-1), 0.5, c_gray_border),
        ('TOPPADDING',   (0,0),(-1,-1), 7),
        ('BOTTOMPADDING',(0,0),(-1,-1), 7),
        ('LEFTPADDING',  (0,0),(-1,-1), 12),
    ]))
    elements.append(summary_t)

    # ── Build ─────────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc_pdf = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=ML, rightMargin=MR,
        topMargin=MT + 1.1*cm, bottomMargin=MB + 0.8*cm,
        title=f'ERS — {project_nome}',
        author=nome_gestor,
        subject='Especificação de Requisitos de Software',
    )
    doc_pdf.build(elements,
                  onFirstPage=_on_first_page,
                  onLaterPages=_on_page)
    return buf.getvalue()


# ── DOCX Builder ──────────────────────────────────────────────────────────────

def _build_docx(project, requirements, topic_ids=None, requirement_ids=None,
                diagramas=None, custom_labels=None):
    filtered = _filter_requirements(requirements, topic_ids, requirement_ids)
    groups   = _group_by_topic(filtered)
    clabels  = custom_labels or {}

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # Title
    title = doc.add_heading(f'ERS — {project.get("nome", "Projeto")}', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(*GREEN)

    # Project info
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gestor = project.get('gestor') or {}
    for line in [
        f'Gestor: {gestor.get("nome", "N/A")}',
        f'Cliente: {project.get("nome_cliente") or "N/A"}',
        f'Data: {now_brt().strftime("%d/%m/%Y")}',
        f'Total de requisitos: {len(filtered)}',
    ]:
        run = info.add_run(line + '\n')
        run.font.size = Pt(10)

    doc.add_paragraph()

    # Description
    section_num = 1
    if project.get('descricao'):
        h = doc.add_heading(f'{section_num}. Descrição do Projeto', level=1)
        h.runs[0].font.color.rgb = RGBColor(*GREEN)
        doc.add_paragraph(project['descricao'])
        section_num += 1

    # Requirements
    for tipo, reqs in groups:
        label = TOPIC_LABELS.get(tipo, clabels.get(tipo, tipo))
        h = doc.add_heading(f'{section_num}. {label}', level=1)
        h.runs[0].font.color.rgb = RGBColor(*GREEN)

        for idx, req in enumerate(reqs, 1):
            codigo = req.get('codigo') or f'{idx:03d}'
            h2 = doc.add_heading(f'{section_num}.{idx}  {codigo} — {req.get("titulo", "")}', level=2)
            if req.get('descricao'):
                doc.add_paragraph(req['descricao'])
            table = doc.add_table(rows=2, cols=2)
            table.style = 'Light Grid Accent 1'
            table.cell(0, 0).text = 'Prioridade'
            table.cell(0, 1).text = PRIORITY_LABELS.get(req.get('prioridade') or 'media', '')
            table.cell(1, 0).text = 'Status'
            table.cell(1, 1).text = STATUS_LABELS.get(req.get('status') or '', '')
            doc.add_paragraph()
        section_num += 1

    # Diagrams
    if diagramas:
        h = doc.add_heading(f'{section_num}. Diagramas', level=1)
        h.runs[0].font.color.rgb = RGBColor(*GREEN)
        for diagrama in diagramas:
            doc.add_heading(diagrama.get('nome', 'Diagrama'), level=2)
            dados = diagrama.get('dados')
            tipo_mime = diagrama.get('tipo_mime', '')
            if dados and 'pdf' not in tipo_mime:
                try:
                    doc.add_picture(io.BytesIO(dados), width=Inches(5.5))
                except Exception:
                    doc.add_paragraph('[Imagem não pôde ser inserida]')
            elif 'pdf' in tipo_mime:
                doc.add_paragraph('[Arquivo PDF — visualize separadamente]')
            doc.add_paragraph()
        section_num += 1

    # Summary
    h = doc.add_heading('Resumo', level=1)
    h.runs[0].font.color.rgb = RGBColor(*GREEN)
    doc.add_paragraph(f'Total de requisitos: {len(filtered)}')
    for tipo, reqs in groups:
        label = TOPIC_LABELS.get(tipo, clabels.get(tipo, tipo))
        doc.add_paragraph(f'{label}: {len(reqs)}')
    if diagramas:
        doc.add_paragraph(f'Diagramas: {len(diagramas)}')

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Fallback minimal PDF ──────────────────────────────────────────────────────

def _build_minimal_pdf(project, requirements, topic_ids=None, requirement_ids=None):
    try:
        from fpdf import FPDF
    except ImportError:
        return _build_raw_pdf(project, requirements, topic_ids, requirement_ids)

    filtered = _filter_requirements(requirements, topic_ids, requirement_ids)
    groups   = _group_by_topic(filtered)

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font('Helvetica', 'B', 18)
    pdf.cell(0, 15, f'ERS - {project.get("nome", "Projeto")}', ln=True, align='C')
    pdf.set_font('Helvetica', '', 10)
    gestor = project.get('gestor') or {}
    pdf.cell(0, 6, f'Gestor: {gestor.get("nome", "N/A")}', ln=True, align='C')
    pdf.cell(0, 6, f'Cliente: {project.get("nome_cliente") or "N/A"}', ln=True, align='C')
    pdf.ln(10)

    if project.get('descricao'):
        pdf.set_font('Helvetica', 'B', 14)
        pdf.cell(0, 10, '1. Descricao do Projeto', ln=True)
        pdf.set_font('Helvetica', '', 10)
        pdf.multi_cell(0, 5, project['descricao'])
        pdf.ln(5)

    section_num = 2
    for tipo, reqs in groups:
        label = TOPIC_LABELS.get(tipo, tipo)
        pdf.set_font('Helvetica', 'B', 14)
        pdf.cell(0, 10, f'{section_num}. {label}', ln=True)
        for idx, req in enumerate(reqs, 1):
            codigo = req.get('codigo') or f'{idx:03d}'
            pdf.set_font('Helvetica', 'B', 11)
            pdf.multi_cell(0, 7, f'{section_num}.{idx} {codigo} - {req.get("titulo", "")}')
            if req.get('descricao'):
                pdf.set_font('Helvetica', '', 10)
                pdf.multi_cell(0, 5, req['descricao'])
            pdf.set_font('Helvetica', '', 9)
            pdf.cell(0, 5,
                     f'Prioridade: {PRIORITY_LABELS.get(req.get("prioridade","media"),"")} | '
                     f'Status: {STATUS_LABELS.get(req.get("status",""),"")}', ln=True)
            pdf.ln(3)
        section_num += 1
    return pdf.output()


def _build_raw_pdf(project, requirements, topic_ids=None, requirement_ids=None):
    filtered = _filter_requirements(requirements, topic_ids, requirement_ids)
    lines = [
        f"ERS - {project.get('nome', 'Projeto')}",
        f"Gestor: {(project.get('gestor') or {}).get('nome', 'N/A')}",
        f"Cliente: {project.get('nome_cliente') or 'N/A'}",
        f"Data: {datetime.now(timezone.utc).strftime('%d/%m/%Y')}",
        "", "AVISO: Instale reportlab para PDF completo.", "",
    ]
    for req in filtered:
        lines.append(f"{req.get('codigo','?')} - {req.get('titulo') or ''}")
        if req.get('descricao'):
            lines.append(f"  {req['descricao']}")
        lines.append("")

    content = "\n".join(lines)
    objects, obj_offsets = [], []

    def add_obj(data):
        objects.append(data)

    add_obj("1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj")
    add_obj("2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj")
    add_obj("3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>\nendobj")
    escaped = content.replace('\\','\\\\').replace('(','\\(').replace(')','\\)')
    sl, y = [], 750
    sl.append("BT\n/F1 10 Tf")
    for line in escaped.split('\n'):
        if y < 50: break
        sl.append(f"1 0 0 1 50 {y} Tm\n({line[:80]}) Tj")
        y -= 14
    sl.append("ET")
    stream = '\n'.join(sl)
    add_obj(f"4 0 obj\n<< /Length {len(stream)} >>\nstream\n{stream}\nendstream\nendobj")
    add_obj("5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj")

    pdf = b"%PDF-1.4\n"
    for obj in objects:
        obj_offsets.append(len(pdf))
        pdf += obj.encode() + b"\n"
    xref = len(pdf)
    pdf += b"xref\n" + f"0 {len(objects)+1}\n".encode() + b"0000000000 65535 f \n"
    for off in obj_offsets:
        pdf += f"{off:010d} 00000 n \n".encode()
    pdf += (f"trailer\n<< /Size {len(objects)+1} /Root 1 0 R >>\n"
            f"startxref\n{xref}\n%%EOF\n").encode()
    return pdf


# ── Public API ────────────────────────────────────────────────────────────────

def generate_ers_pdf(project, requirements, topic_ids=None, requirement_ids=None,
                     diagramas=None, custom_labels=None):
    project_name = project.get('nome', 'projeto').replace(' ', '_')
    return f'ERS_{project_name}.pdf', _build_pdf(
        project, requirements, topic_ids, requirement_ids,
        diagramas=diagramas, custom_labels=custom_labels)


def generate_ers_docx(project, requirements, topic_ids=None, requirement_ids=None,
                      diagramas=None, custom_labels=None):
    project_name = project.get('nome', 'projeto').replace(' ', '_')
    return f'ERS_{project_name}.docx', _build_docx(
        project, requirements, topic_ids, requirement_ids,
        diagramas=diagramas, custom_labels=custom_labels)
